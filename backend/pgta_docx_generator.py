"""
DOCX Report Generator for PGT-A Reports
Generates Word documents matching the PDF template with 1:1 precision.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pgta_classify as clf
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
import os
import sys
from io import BytesIO
from datetime import datetime

def set_cell_border(cell, **kwargs):
    """
    Set cell's border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "single", "color": "#0000FF"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existence, if none found, then create one
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for edge in ('start', 'top', 'end', 'bottom', 'left', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existence, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


class PGTADocxGenerator:
    """Generates DOCX reports for PGT-A with pixel-level precision"""
    
    # Static content (same as PDF template)
    METHODOLOGY_TEXT = """Chromosomal aneuploidy analysis was performed using ChromInst® PGT-A kit from Yikon Genomics (Suzhou) Co., Ltd - China. The Yikon - ChromInst® PGT-A kit with the Genemind - SURFSeq 5000* High-throughput Sequencing Platform allows detection of aneuploidies in all 23 sets of Chromosomes. Probes are not covering the p arm of acrocentric chromosomes as they are rich in repeat regions and RNA markers and devoid of genes. Changes in this region will not be detected. However, these regions have less clinical significance due to the absence of genes. Chromosomal aneuploidy can be detected by copy number variations (CNVs), which represent a class of variation in which segments of the genome have been duplicated (gains) or deleted (losses). Large, genomic copy number imbalances can range from sub-chromosomal regions to entire chromosomes. Inherited and de-novo CNVs (up to 10 Mb) have been associated with many disease conditions. This assay was performed on DNA extracted from embryo biopsy samples."""
    
    MOSAICISM_TEXT = """Mosaicism arises in the embryo due to mitotic errors which lead to the production of karyotypically distinct cell lineages within a single embryo [1]. NGS has the sensitivity to detect mosaicism when 30% or the above cells are abnormal [2]. Mosaicism is reported in our laboratory as follows [3]."""
    
    MOSAICISM_BULLETS = [
        "Embryos with less than 30% mosaicism are considered as euploid.",
        "Embryos with 30% to 50% mosaicism will be reported as low level mosaic, 51% to 80% mosaicism will be reported as high level mosaic.",
        "When three chromosomes or more than three chromosomes showing mosaic change, it will be denoted as complex mosaic.",
        "If greater than 80% mosaicism detected in an embryo it will be considered aneuploid."
    ]
    
    MOSAICISM_CLINICAL = """Clinical significance of transferring mosaic embryos is still under evaluation. Based on Preimplantation Genetic Diagnosis International Society (PGDIS) Position Statement – 2019 transfer of these embryos should be considered only after appropriate counselling of the patient and alternatives have been discussed. Invasive prenatal testing with karyotyping in the amniotic fluid needs to be advised in such cases [4]. As shown in published literature evidence, such transfers can result in normal pregnancy or miscarriage or an offspring with chromosomal mosaicism [5,6,7]."""
    
    LIMITATIONS = [
        "This technique cannot detect point mutations, balanced translocations, inversions, triploidy, uniparental disomy and epigenetic modifications.",
        "Probes used do not cover the p arm of acrocentric chromosomes as they are rich in repeat regions and RNA markers and devoid of genes. Changes in this region will not be detected. However, these regions have less clinical significance due to the absence of genes.",
        "Deletions and duplications with the size of < 10 Mb cannot be detected.",
        "Risk of misinterpretation of the actual embryo karyotype due to the presence of chromosomal mosaicism, either at cleavage-stage or at blastocyst stage may exist.",
        "This technique cannot detect variants of polyploidy and haploidy",
        "NGS without genotyping cannot identify the nature (meiotic or mitotic) nor the parental origin of aneuploidies",
        "Due to the intrinsic nature of chromosomal mosaicism, the chromosomal make-up achieved from a biopsy only may represent a picture of a small part of the embryo and may not necessarily reflect the chromosomal content of the entire embryo. Also, the mosaicism level inferred from a multi-cell TE biopsy might not unequivocally represent the exact chromosomal mosaicism percentage of the TE cells or the inner cell mass constitution."
    ]
    
    REFERENCES = [
        'McCoy, Rajiv C. "Mosaicism in Preimplantation human embryos: when chromosomal abnormalities are the norm." Trends in genetics 33.7 (2017): 448-463.',
        'ESHRE PGT-SR/PGT-A Working Group, et al. "ESHRE PGT Consortium good practice recommendations for the detection of structural and numerical chromosomal aberrations." Human reproduction open 2020.3 (2020): hoaa017.',
        'ESHRE Working Group on Chromosomal Mosaicism, et al. "ESHRE survey results and good practice recommendations on managing chromosomal mosaicism." Hum Reprod Open. 2022 Nov 7;2022(4):hoac044.',
        'Cram, D. S., et al. "PGDIS position statement on the transfer of mosaic embryos 2019." Reproductive biomedicine online 39 (2019): e1-e4.',
        'Victor, Andrea R., et al. "One hundred mosaic embryos transferred prospectively in a single clinic: exploring when and why they result in healthy pregnancies." Fertility and sterility 111.2 (2019): 280-293.',
        'Lin, Pin-Yao, et al. "Clinical outcomes of single mosaic embryo transfer: high-level or low-level mosaic embryo, does it matter?" Journal of clinical medicine 9.6 (2020): 1695.',
        'Kahraman, Semra, et al. "The birth of a baby with mosaicism resulting from a known mosaic embryo transfer: a case report." Human Reproduction 35.3 (2020): 727-733.'
    ]

    def __init__(self, assets_dir="assets/pgta"):
        """Initialize assets and log paths"""
        script_dir = os.path.dirname(os.path.abspath(__file__))
        self.assets_dir = os.path.join(script_dir, assets_dir)
        
        self.header_logo = os.path.join(self.assets_dir, "image_page1_0.png")
        self.footer_banner = os.path.join(self.assets_dir, "image_page1_1.png")
        self.genqa_logo = os.path.join(self.assets_dir, "genqa_logo.png")
        self.signs_image = os.path.join(self.assets_dir, "signs.png")

    # --- OXML PRECISION HELPERS ---
    
    def _set_cell_background(self, cell, fill):
        """Set background shading for a table cell using OXML"""
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), fill.replace('#', ''))
        cell._tc.get_or_add_tcPr().append(shading_elm)

    def _set_table_fixed_layout(self, table):
        """Force a table to use a fixed layout so column widths are strictly respected"""
        tbl_pr = table._element.xpath('w:tblPr')[0]
        layout = OxmlElement('w:tblLayout')
        layout.set(qn('w:type'), 'fixed')
        tbl_pr.append(layout)

    def _set_column_widths(self, table, widths_pt):
        """Set exact column widths in points (1 pt = 1/72 inch) for every row/cell"""
        total_width = sum(widths_pt)
        # Set total table width
        tbl_pr = table._element.xpath('w:tblPr')[0]
        tbl_w = OxmlElement('w:tblW')
        tbl_w.set(qn('w:w'), str(int(total_width * 20))) 
        tbl_w.set(qn('w:type'), 'dxa')
        tbl_pr.append(tbl_w)

        # Iterate rows and set each cell's width to ensure parity even if columns[] fails
        for row in table.rows:
            for i, width in enumerate(widths_pt):
                if i < len(row.cells):
                    cell = row.cells[i]
                    tc_pr = cell._tc.get_or_add_tcPr()
                    tc_w = OxmlElement('w:tcW')
                    tc_w.set(qn('w:w'), str(int(width * 20))) 
                    tc_w.set(qn('w:type'), 'dxa')
                    tc_pr.append(tc_w)

    def _set_paragraph_font(self, paragraph, font_name="Segoe UI", font_size=9, bold=False, italic=False, color=None):
        """Apply font styling to every run in a paragraph to ensure 1:1 PDF parity"""
        if not paragraph.runs:
            paragraph.add_run()
        for run in paragraph.runs:
            run.font.name = font_name
            r = run._element
            r.get_or_add_rPr().get_or_add_rFonts().set(qn('w:ascii'), font_name)
            r.get_or_add_rPr().get_or_add_rFonts().set(qn('w:hAnsi'), font_name)
            
            run.font.size = Pt(font_size)
            run.bold = bold
            run.italic = italic
            if color:
                if isinstance(color, str) and color.startswith('#'):
                    run.font.color.rgb = RGBColor.from_string(color[1:])
                else:
                    run.font.color.rgb = color

    def _clean(self, val, default=""):
        """Sanitize values"""
        if val is None: return default
        s = str(val).strip()
        if s.lower() == "nan": return default
        return s

    # --- GENERATION LOGIC ---

    def generate_docx(self, output_path, patient_data, embryos_data, show_logo=True, show_grid=False):
        """Main entry point for DOCX generation"""
        self.show_grid = show_grid
        doc = Document()
        
        # 1. Page Setup (Margins mirroring PDF exactly)
        sections = doc.sections
        # 1. Page Setup (US Letter: 612pt x 792pt)
        sections = doc.sections
        for section in sections:
            section.page_width = Pt(612)
            section.page_height = Pt(792)
            section.top_margin = Pt(70)
            section.bottom_margin = Pt(60)
            section.left_margin = Pt(58)
            section.right_margin = Pt(58)
            section.header_distance = Pt(20)
            section.footer_distance = Pt(20)
        
        # Global Font Defaults
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(9)
        
        # 2. Cover Page
        self._add_cover_page(doc, patient_data, embryos_data)
        
        # 3. Headers/Footers
        self._setup_page_header_footer(doc, show_logo=show_logo)
        
        # 4. Methodology Page (pass embryos so PGDIS note is conditional)
        doc.add_page_break()
        self._add_methodology_page(doc, embryos_data)

        # 5. Check if ALL embryos are "Low DNA concentration" — mirrors PDF logic exactly
        all_low_dna = bool(embryos_data)
        for embryo in embryos_data:
            interp = str(embryo.get('interpretation', '')).upper()
            res    = str(embryo.get('result_summary', '')).upper()
            if "LOW DNA" not in interp and "LOW DNA" not in res:
                all_low_dna = False
                break

        if all_low_dna:
            # All Low DNA: append signature right after methodology, no embryo pages
            doc.add_paragraph()
            self._add_signature_section(doc)
        else:
            doc.add_page_break()
            # 6. Individual embryo pages — skip Low DNA embryos, add signature after each
            for embryo in embryos_data:
                interp = str(embryo.get('interpretation', '')).upper()
                res    = str(embryo.get('result_summary', '')).upper()
                if "LOW DNA" in interp or "LOW DNA" in res:
                    continue
                self._add_embryo_page(doc, patient_data, embryo)

        # 7. Save
        doc.save(output_path)
        return output_path

    def _setup_page_header_footer(self, doc, show_logo=True):
        """Setup branding in headers and footers using locked table layouts"""
        for section in doc.sections:
            # Header
            header = section.header
            header.paragraphs[0].clear()
            if show_logo and self.header_logo and os.path.exists(self.header_logo):
                p = header.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(self.header_logo, width=Pt(496))

            # Footer
            footer = section.footer
            footer.paragraphs[0].clear()
            footer_table = footer.add_table(rows=1, cols=2, width=Pt(496))
            self._set_table_fixed_layout(footer_table)
            self._set_column_widths(footer_table, [416, 80])
            
            # Banner
            if show_logo and self.footer_banner and os.path.exists(self.footer_banner):
                c0 = footer_table.rows[0].cells[0]
                p0 = c0.paragraphs[0]
                p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p0.add_run().add_picture(self.footer_banner, width=Pt(416))
            
            # GenQA
            if self.genqa_logo and os.path.exists(self.genqa_logo):
                c1 = footer_table.rows[0].cells[1]
                p1 = c1.paragraphs[0]
                p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                p1.add_run().add_picture(self.genqa_logo, width=Pt(65))

    def _add_cover_page(self, doc, patient_data, embryos_data):
        """Cover page mirroring PDF layout and colors"""
        # Title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("Preimplantation Genetic Testing for Aneuploidies (PGT-A)")
        self._set_paragraph_font(title, font_name="Calibri", font_size=14, bold=False)
        
        doc.add_paragraph() # Spacer
        
        # Patient Info Table [108, 12, 131, 108, 12, 119] - 6 rows (spouse name combined with patient name)
        # Adjusted widths to give more space to label columns to prevent date wrap
        info_table = doc.add_table(rows=6, cols=6)
        self._apply_grid_to_table(info_table)
        self._set_table_fixed_layout(info_table)
        self._set_column_widths(info_table, [108, 12, 131, 108, 12, 119])
        self._populate_patient_table(info_table, patient_data, is_embryo=False)
        
        doc.add_paragraph() # Spacer
        
        # PNDT Disclaimer — bold + italic (PNDT Act 1994)
        disclaimer = doc.add_paragraph()
        disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_d = disclaimer.add_run(
            "This test does not reveal sex of the fetus & confers to PNDT act, 1994"
        )
        run_d.bold   = True
        run_d.italic = True
        run_d.font.size = Pt(9.5)
        
        doc.add_paragraph() # Spacer
        
        # Indication
        if 'indication' in patient_data and patient_data['indication']:
            p_ind = doc.add_paragraph()
            self._set_paragraph_font(p_ind, font_name="Calibri", font_size=10, bold=True)
            p_ind.add_run("Indication")
            p_val = doc.add_paragraph(self._clean(patient_data['indication']))
            self._set_paragraph_font(p_val, font_size=9)
            doc.add_paragraph()

        # Results Summary Header
        p_res = doc.add_paragraph()
        self._set_paragraph_font(p_res, font_name="Calibri", font_size=10, bold=True)
        p_res.add_run("Results summary")
        
        # Results Summary Table [50, 95, 185, 80, 86]
        res_table = doc.add_table(rows=len(embryos_data) + 1, cols=5)
        self._apply_grid_to_table(res_table)
        self._set_table_fixed_layout(res_table)
        self._set_column_widths(res_table, [50, 95, 185, 80, 86])
        
        # Row 0: Headers (Peach bg)
        headers = ['S. No.', 'Sample', 'Result', 'MTcopy', 'Interpretation']
        for i, h in enumerate(headers):
            cell = res_table.rows[0].cells[i]
            cell.text = h
            self._set_paragraph_font(cell.paragraphs[0], font_size=9, bold=True)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._set_cell_background(cell, "F9BE8F")

        # Data rows (F1F1F7 bg)
        for i, emb in enumerate(embryos_data, 1):
            row = res_table.rows[i]
            row.cells[0].text = str(i)
            row.cells[1].text = self._clean(emb.get('embryo_id'))

            raw = self._clean(emb.get('result_summary') or emb.get('result_description') or '')
            info = clf.classify_embryo(raw)

            # Result → mapped display text; Interpretation → always "NA"
            res_display   = info["summary_text"]
            interp_display= "NA"
            cell_color    = info["is_abnormal"] and "#FF0000" or "#000000"

            # MTcopy: "NA" unless explicitly provided
            raw_mt = self._clean(emb.get('mtcopy', ''))
            mt = raw_mt if raw_mt and raw_mt.upper() not in ('NA', 'N/A', '') else "NA"

            row.cells[2].text = res_display
            row.cells[3].text = mt
            row.cells[4].text = interp_display

            for c_idx, cell in enumerate(row.cells):
                self._set_cell_background(cell, "F1F1F7")
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Color Result and Interpretation columns
                color = cell_color if c_idx in (2, 4) else None
                self._set_paragraph_font(p, font_size=9, color=color)
        
        # Results Summary Comment (optional, appears below table)
        results_summary_comment = self._clean(patient_data.get('results_summary_comment', ''))
        if results_summary_comment:
            doc.add_paragraph()  # Spacer
            p_comment = doc.add_paragraph(results_summary_comment)
            self._set_paragraph_font(p_comment, font_size=9)

    def _populate_patient_table(self, table, data, is_embryo=False):
        """Standard Patient Info Population Logic"""
        # Patient name and spouse name - spouse on new line
        import re
        patient_name = re.sub(r'\s+', ' ', self._clean(data.get('patient_name'))).strip()
        spouse_name = re.sub(r'\s+', ' ', self._clean(data.get('spouse_name'))).strip()
        # Put spouse on new line if present
        combined_name = f"{patient_name}\n{spouse_name}" if spouse_name else patient_name
        
        rows_map = [
            ("PATIENT NAME", combined_name, "PIN", "pin"),
            ("DATE OF BIRTH/ AGE", "age", "SAMPLE NUMBER", "sample_number"),
            ("REFERRING CLINICIAN", "referring_clinician", "BIOPSY DATE", "biopsy_date"),
            ("HOSPITAL/CLINIC", "hospital_clinic", "SAMPLE COLLECTION DATE", "sample_collection_date"),
            ("SPECIMEN", "specimen", "SAMPLE RECEIPT DATE", "sample_receipt_date"),
            ("BIOPSY PERFORMED BY", "biopsy_performed_by", "REPORT DATE", "report_date")
        ]
        for r_idx, (l1, v1, l2, v2) in enumerate(rows_map):
            if r_idx >= len(table.rows): break
            row = table.rows[r_idx]
            
            # Populate labels and colons
            if l1: row.cells[0].text = l1; row.cells[1].text = ":"
            if l2: row.cells[3].text = l2; row.cells[4].text = ":"
            
            # Populate cleaned values - first row has combined name directly
            if v1: 
                if r_idx == 0:  # First row - combined name already a string
                    row.cells[2].text = v1
                else:
                    row.cells[2].text = self._clean(data.get(v1))
            if v2: row.cells[5].text = self._clean(data.get(v2))
            
            for cell_idx, cell in enumerate(row.cells):
                self._set_cell_background(cell, "F1F1F7")
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                self._set_paragraph_font(cell.paragraphs[0], font_name="Segoe UI", font_size=10, bold=True)
                p_fmt = cell.paragraphs[0].paragraph_format
                p_fmt.space_before = Pt(2)
                p_fmt.space_after = Pt(2)
                
                # Set cell alignment to match PDF strictly left aligned for values
                # Fixed: Use cell_idx from enumerate instead of row.cells.index(cell) to avoid tuple.index error
                if cell_idx in [0, 3]:  # Label columns
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                    
                    # Logic: PIN label in embryo banner should be right-aligned (flushed to colon)
                    # Page 1 and other labels should remain left-aligned with 12pt padding
                    label_text = cell.text.strip()
                    if is_embryo and label_text == "PIN":
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        cell.paragraphs[0].paragraph_format.left_indent = Pt(0)
                        cell.paragraphs[0].paragraph_format.right_indent = Pt(12)
                    else:
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                        cell.paragraphs[0].paragraph_format.left_indent = Pt(4)
                elif cell_idx in [1, 4]:  # Colon columns
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:  # Value columns
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    def _add_methodology_page(self, doc, embryos_data=None):
        """Methods, Limitations, and References with natural flow but orphan protection"""
        # PGDIS 2019 counselling note only when a mosaic embryo exists
        mosaicism_clinical_item = (None, self.MOSAICISM_CLINICAL, None) if clf.any_mosaic(embryos_data or []) else None

        sections = [
            ("Methodology", self.METHODOLOGY_TEXT, None),
            ("Conditions for reporting mosaicism", self.MOSAICISM_TEXT, self.MOSAICISM_BULLETS),
        ]
        if mosaicism_clinical_item:
            sections.append(mosaicism_clinical_item)
        sections += [
            ("Limitations", None, self.LIMITATIONS),
            ("References", None, [f"{i}. {r}" for i, r in enumerate(self.REFERENCES, 1)])
        ]
        
        for head, body, bullets in sections:
            if head:
                p = doc.add_paragraph()
                self._set_paragraph_font(p, font_size=11, bold=True)
                p.add_run(head)
                # Only keep with next if there's content following
                if body or bullets:
                    p.paragraph_format.keep_with_next = True
            
            if body:
                p = doc.add_paragraph(body)
                self._set_paragraph_font(p, font_size=9)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                # Only keep with bullets if they exist
                if bullets:
                    p.paragraph_format.keep_with_next = True
                
            if bullets:
                for i, b in enumerate(bullets):
                    p = doc.add_paragraph(b, style='List Bullet')
                    self._set_paragraph_font(p, font_size=9)
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    # Keep bullets together naturally
                    if i < len(bullets) - 1:
                        p.paragraph_format.keep_with_next = True
            doc.add_paragraph()

    def _add_embryo_page(self, doc, patient_data, embryo_data):
        """Individual Embryo Result Page with exact PDF metrics"""
        doc.add_page_break()

        # Title repeated on every embryo page — mirrors PDF _build_embryo_page
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._set_paragraph_font(p_title, font_name="Calibri", font_size=14, bold=False)
        p_title.add_run("Preimplantation Genetic Testing for Aneuploidies (PGT-A)")
        doc.add_paragraph()

        # 1. Banner [Total: 490pt] - Match exact cover page positioning
        banner = doc.add_table(rows=2, cols=6)
        self._apply_grid_to_table(banner)
        self._set_table_fixed_layout(banner)
        # Optimized layout: Push PIN block right. PATIENT NAME (82), Colons (12x2), PIN label (24).
        self._set_column_widths(banner, [82, 12, 242, 24, 12, 118])
        # Ensure table is aligned to the left like cover page
        banner.alignment = WD_ALIGN_PARAGRAPH.LEFT
        self._populate_patient_table(banner, patient_data, is_embryo=True)

        doc.add_paragraph()

        # PNDT disclaimer on every embryo page — bold + italic (PNDT Act 1994)
        p_pndt = doc.add_paragraph()
        p_pndt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_pndt = p_pndt.add_run(
            "This test does not reveal sex of the fetus & confers to PNDT act, 1994"
        )
        run_pndt.bold   = True
        run_pndt.italic = True
        run_pndt.font.size = Pt(9.5)

        doc.add_paragraph()

        # 2. Embryo ID - Use embryo_id_detail for detail pages, fallback to embryo_id
        eid = self._clean(embryo_data.get('embryo_id_detail')) or self._clean(embryo_data.get('embryo_id'))
        p_eid = doc.add_paragraph()
        self._set_paragraph_font(p_eid, font_name="Calibri", font_size=12, bold=True, color="#1F497D")
        p_eid.add_run(f"EMBRYO: {eid}")
        
        # 3. Summary [Total: 490pt]
        raw_result = self._clean(embryo_data.get('result_summary') or embryo_data.get('result_description') or '')
        info = clf.classify_embryo(raw_result)

        # Result: mapped full sentence
        res = info["result_text"]

        # Autosomes: "Normal" if euploid, else auto-derive
        existing_auto = self._clean(embryo_data.get('autosomes', ''))
        chr_statuses = embryo_data.get('chromosome_statuses') or {}
        if not chr_statuses:
            chr_statuses = clf.derive_chromosome_statuses(raw_result)
            chr_statuses = clf.validate_statuses(chr_statuses, raw_result)
        auto = clf.derive_autosomes(raw_result, chr_statuses, existing_auto)

        # Sex chromosomes: sanitise, never XX/XY
        existing_sex = self._clean(embryo_data.get('sex_chromosomes', ''))
        sex = clf.sanitize_sex_chromosomes(existing_sex, raw_result, info["classification"])

        # Interpretation: always "NA"
        interp = "NA"

        # MTcopy: "NA" unless explicitly provided
        raw_mt = self._clean(embryo_data.get('mtcopy', ''))
        mt = raw_mt if raw_mt and raw_mt.upper() not in ('NA', 'N/A', '') else "NA"

        cell_color   = "#FF0000" if info["is_abnormal"] else "#000000"
        sex_color    = cell_color if sex.upper() not in ('NORMAL', '') else "#000000"
        details = [
            ("Result:", res, "#000000"),
            ("Autosomes:", auto, cell_color),
            ("Sex Chromosomes:", sex, sex_color),
            ("Interpretation:", interp, cell_color),
            ("MTcopy:", mt, "#000000"),
        ]
        
        d_table = doc.add_table(rows=len(details), cols=1)
        self._apply_grid_to_table(d_table)
        self._set_table_fixed_layout(d_table)
        self._set_column_widths(d_table, [490])
        for idx, (label, val, color) in enumerate(details):
            cell = d_table.rows[idx].cells[0]
            self._set_cell_background(cell, "F1F1F7")
            p = cell.paragraphs[0]
            self._set_paragraph_font(p, font_size=9, bold=True)
            p.add_run(f"{label} ")
            run_val = p.add_run(val)
            self._set_paragraph_font(p, font_size=9, bold=False, color=color)
            p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)

        doc.add_paragraph()
        
        # 4. Chart
        p_ch = doc.add_paragraph()
        self._set_paragraph_font(p_ch, font_size=10, bold=True)
        p_ch.add_run("COPY NUMBER CHART")
        if embryo_data.get('cnv_image_path') and os.path.exists(embryo_data['cnv_image_path']):
            doc.add_picture(embryo_data['cnv_image_path'], width=Pt(496))
        
        doc.add_paragraph()
        
        # 5. CNV Status Table [Total: 496pt] - Skip for Inconclusive results
        result_summary = self._clean(embryo_data.get('result_summary', ''))
        result_desc = self._clean(embryo_data.get('result_description', ''))
        is_inconclusive = "INCONCLUSIVE" in result_summary.upper() or "INCONCLUSIVE" in result_desc.upper() or "INCONCLUSIVE" in interp.upper()
        
        # Add inconclusive comment under CNV chart if present
        if is_inconclusive:
            inconclusive_comment = self._clean(embryo_data.get('inconclusive_comment', ''))
            if inconclusive_comment:
                comment_p = doc.add_paragraph(inconclusive_comment)
                self._set_paragraph_font(comment_p, font_size=11)
        
        if not is_inconclusive:
            chr_statuses = embryo_data.get('chromosome_statuses') or {}
            if not chr_statuses:
                chr_statuses = clf.derive_chromosome_statuses(raw_result)
                chr_statuses = clf.validate_statuses(chr_statuses, raw_result)
            mosaic_map = embryo_data.get('mosaic_percentages', {})
            
            autosomes = str(embryo_data.get('autosomes', '')).upper()
            sex_chrs = str(embryo_data.get('sex_chromosomes', '')).upper()
            
            import re as re_mos
            has_mosaic = any(
                v and str(v).strip() and str(v).strip() != '-' and re_mos.search(r'\d', str(v))
                for v in mosaic_map.values()
            )
            
            is_autosomes_normal = 'NORMAL' in autosomes or 'EUPLOID' in autosomes or not autosomes.strip()
            is_sex_mosaic = 'MOSAIC' in sex_chrs
            
            if is_autosomes_normal and is_sex_mosaic:
                has_mosaic = False
                
            num_rows = 3 if has_mosaic else 2
            cnv_table = doc.add_table(rows=num_rows, cols=23)
            self._apply_grid_to_table(cnv_table)
            self._set_table_fixed_layout(cnv_table)
            self._set_column_widths(cnv_table, [75] + [19.13]*22)
            
            # Header Row
            cnv_table.rows[0].cells[0].text = "Chromosome"
            for i in range(1, 23): cnv_table.rows[0].cells[i].text = str(i)
            
            # Status Row
            cnv_table.rows[1].cells[0].text = "CNV status"
            for i in range(1, 23):
                cell = cnv_table.rows[1].cells[i]
                stat = str(chr_statuses.get(str(i), 'N'))
                cell.text = stat
                color = self._get_status_color_docx(stat)
                self._set_paragraph_font(cell.paragraphs[0], font_size=8, bold=True, color=color)
                
            # Mosaic Row
            if has_mosaic:
                cnv_table.rows[2].cells[0].text = "Mosaic (%)"
                for i in range(1, 23):
                    cnv_table.rows[2].cells[i].text = str(mosaic_map.get(str(i), '-'))

            for row in cnv_table.rows:
                for c_idx, cell in enumerate(row.cells):
                    self._set_cell_background(cell, "F1F1F7")
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if c_idx == 0: p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    self._set_paragraph_font(p, font_size=8, bold=True)

            # CNV legend — mirrors PDF exactly
            legend_p = doc.add_paragraph()
            legend_run = legend_p.add_run(
                "N \u2013 Normal, G-Gain, L-Loss, SG-Segmental Gain, SL-Segmental Loss, "
                "M-Mosaic, MG- Mosaic Gain, ML-Mosaic Loss, SMG-Segmental Mosaic Gain, "
                "SML-Segmental Mosaic Loss"
            )
            legend_run.italic = True
            legend_run.font.size = Pt(7)

        doc.add_paragraph()
        self._add_signature_section(doc)
        # Page break after each embryo+signature block — mirrors PDF PageBreak()
        doc.add_page_break()

    def _add_signature_section(self, doc):
        """Pixel-Perfect 3-Column Signature Section"""
        table = doc.add_table(rows=2, cols=3)
        self._set_table_fixed_layout(table)
        self._set_column_widths(table, [156, 156, 156])
        
        if self.signs_image and os.path.exists(self.signs_image):
            p = table.rows[0].cells[1].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(self.signs_image, width=Pt(300))
        
        sigs = [
            ("Anand Babu. K, Ph.D",        "Molecular Biologist"),
            ("Sachin D Honguntikar, Ph.D",  "Molecular Geneticist"),
            ("Dr Suriyakumar G",            "Director"),
        ]
        for i, (name, title) in enumerate(sigs):
            cell = table.rows[1].cells[i]
            p1 = cell.paragraphs[0]; p1.text = name; p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._set_paragraph_font(p1, font_size=11)
            p2 = cell.add_paragraph(title); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._set_paragraph_font(p2, font_size=11)

    def _get_result_color_hex(self, res, interp=None):
        """Return red if embryo is abnormal/mosaic, black if normal."""
        combined = " ".join(filter(None, [str(res or ""), str(interp or "")]))
        return "#FF0000" if clf.classify_embryo(combined)["is_abnormal"] else "#000000"

    def _get_status_color_docx(self, status):
        """CNV status table color: red for L/G/SL/SG, orange for mosaic, grey for NR, black for N."""
        s = str(status).upper().strip()
        if not s or s == 'N':
            return "#000000"
        if s in ('NR', 'FAILED', 'NO RESULT'):
            return "#808080"
        if s in ('ML', 'MG', 'SML', 'SMG', 'M', 'SML/SMG', 'SMG/SML'):
            return "#FF8C00"
        if s in ('L', 'G', 'SL', 'SG', 'SL/SG', 'SG/SL'):
            return "#FF0000"
        # Numeric mosaic percentage
        try:
            if float(s.replace('%', '')) > 0:
                return "#FF8C00"
        except Exception:
            pass
        return "#000000"

    def _apply_grid_to_table(self, table):
        """Apply lite white grid lines to table if enabled"""
        if not hasattr(self, 'show_grid') or not self.show_grid:
            return
            
        grid_color = "E0E0E0" # Lite white/grey
        for row in table.rows:
            for cell in row.cells:
                set_cell_border(
                    cell,
                    top={"sz": 4, "val": "single", "color": grid_color},
                    bottom={"sz": 4, "val": "single", "color": grid_color},
                    start={"sz": 4, "val": "single", "color": grid_color},
                    end={"sz": 4, "val": "single", "color": grid_color}
                )
